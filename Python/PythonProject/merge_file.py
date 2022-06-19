


import os
from docx import Document
from docxcompose.composer import Composer

files = ['Assignment_1 .docx', 'Assignment_1.docx', 'Assignment_2 .docx', 'Assignment_2.docx', 'Assignment_3 .docx', 'Assignment_3.docx', 'Assignment_4 .docx', 'Assignment_4.docx', 'Assignment_5.docx', 'Programming_Assingment1.docx', 'Programming_Assingment2 (1).docx', 'Programming_Assingment2.docx', 'Programming_Assingment20.docx', 'Programming_Assingment3 (1).docx', 'Programming_Assingment3.docx', 'Programming_Assingment4 (1).docx', 'Programming_Assingment4.docx']
composed = "empty.docx"


result = Document(files[0])
# result.add_page_break()
composer = Composer(result)

for i in range(0, len(files)):
    doc = Document(files[i])

    # if i != len(files) - 1:
    #     doc.add_page_break()

    composer.append(doc)

composer.save(composed)
os.system(composed)


# from docx import Document
# from docxcompose.composer import Composer
    
# master = Document("out.docx")
# composer = Composer(master)
# doc1 = Document("in.docx")
# composer.append(doc1)
# master.save('out.docx') 